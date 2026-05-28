import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
import threading
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF
from PIL import Image
import io
import tempfile
import time


class SimplePDFConverter:
    def __init__(self, root):
        self.root = root
        root.title("PDF转Word工具")
        root.geometry("500x400")

        # 文件夹路径变量
        self.folder_path = tk.StringVar()

        # 创建界面
        self.create_widgets()

    def create_widgets(self):
        """创建界面组件"""
        # 标题
        tk.Label(self.root, text="PDF批量转Word工具",
                 font=("微软雅黑", 16, "bold")).pack(pady=20)

        # 文件夹选择框架
        frame = tk.Frame(self.root)
        frame.pack(pady=10, padx=20, fill='x')

        tk.Label(frame, text="文件夹:",
                 font=("微软雅黑", 11)).pack(side='left')

        # 文件夹显示框
        entry = tk.Entry(frame, textvariable=self.folder_path,
                         font=("微软雅黑", 10), width=30)
        entry.pack(side='left', padx=10)

        # 浏览按钮
        browse_btn = tk.Button(frame, text="浏览",
                               command=self.browse_folder,
                               font=("微软雅黑", 10))
        browse_btn.pack(side='left')

        # 处理按钮
        self.process_btn = tk.Button(self.root, text="开始处理",
                                     command=self.start_process,
                                     font=("微软雅黑", 12),
                                     bg="#4CAF50", fg="white",
                                     padx=20, pady=5)
        self.process_btn.pack(pady=20)

        # 进度条
        self.progress = ttk.Progressbar(self.root, mode='indeterminate', length=400)
        self.progress.pack(pady=10)

        # 状态标签
        self.status_label = tk.Label(self.root, text="",
                                     font=("微软雅黑", 10))
        self.status_label.pack()

        # 日志文本框
        tk.Label(self.root, text="处理日志:",
                 font=("微软雅黑", 11)).pack(anchor='w', padx=20, pady=(20, 5))

        self.log_text = tk.Text(self.root, height=8, width=50,
                                font=("Consolas", 9))
        self.log_text.pack(padx=20, fill='both', expand=True)

    def browse_folder(self):
        """浏览文件夹"""
        folder = filedialog.askdirectory(title="选择文件夹")
        if folder:
            self.folder_path.set(folder)
            self.log(f"已选择文件夹: {folder}")

    def log(self, message):
        """添加日志"""
        timestamp = time.strftime("%H:%M:%S")
        self.log_text.insert(tk.END, f"[{timestamp}] {message}\n")
        self.log_text.see(tk.END)

    def start_process(self):
        """开始处理"""
        folder = self.folder_path.get()

        if not folder or not os.path.exists(folder):
            messagebox.showerror("错误", "请选择有效的文件夹！")
            return

        # 禁用按钮
        self.process_btn.config(state='disabled', text="处理中...")

        # 启动进度条
        self.progress.start()
        self.status_label.config(text="正在处理，请稍候...")

        # 在新线程中运行
        thread = threading.Thread(target=self.convert_pdfs, args=(folder,))
        thread.daemon = True
        thread.start()

        # 检查线程状态
        self.check_thread(thread)

    def convert_pdfs(self, folder):
        """转换PDF文件"""
        try:
            self.log("开始扫描PDF文件...")

            # 查找所有PDF文件
            pdf_files = []
            for root, dirs, files in os.walk(folder):
                for file in files:
                    if file.lower().endswith('.pdf'):
                        pdf_files.append(os.path.join(root, file))

            if not pdf_files:
                self.log("未找到PDF文件！")
                self.root.after(0, self.process_complete, False, "未找到PDF文件")
                return

            self.log(f"找到 {len(pdf_files)} 个PDF文件")

            # 创建Word文档
            doc = Document()

            # 设置页面
            section = doc.sections[0]
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

            # 添加标题
            doc.add_heading('PDF转换结果', 0)

            # 处理每个PDF文件
            for i, pdf_file in enumerate(pdf_files, 1):
                self.log(f"处理文件 {i}/{len(pdf_files)}: {os.path.basename(pdf_file)}")

                # 添加文件名称
                doc.add_heading(os.path.basename(pdf_file), level=1)

                # 转换为图片并添加到文档
                images = self.pdf_to_images(pdf_file)
                for img in images:
                    # 保存为临时文件
                    temp_file = tempfile.NamedTemporaryFile(suffix='.jpg', delete=False)
                    img.save(temp_file.name, 'JPEG', quality=90)

                    # 添加图片
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(temp_file.name, width=Inches(6))

                    # 清理临时文件
                    os.unlink(temp_file.name)

                doc.add_page_break()

            # 保存文档
            output_path = os.path.join(folder, f"PDF转换_{time.strftime('%Y%m%d_%H%M%S')}.docx")
            doc.save(output_path)

            self.log(f"文档已保存: {output_path}")
            self.root.after(0, self.process_complete, True, output_path)

        except Exception as e:
            self.log(f"处理出错: {str(e)}")
            self.root.after(0, self.process_complete, False, str(e))

    def pdf_to_images(self, pdf_path):
        """PDF转图片"""
        images = []

        try:
            doc = fitz.open(pdf_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                mat = fitz.Matrix(2.0, 2.0)
                pix = page.get_pixmap(matrix=mat)

                img_data = pix.tobytes("ppm")
                img = Image.open(io.BytesIO(img_data))
                images.append(img)

            doc.close()
        except Exception as e:
            self.log(f"转换失败 {pdf_path}: {e}")

        return images

    def check_thread(self, thread):
        """检查线程状态"""
        if thread.is_alive():
            self.root.after(100, self.check_thread, thread)

    def process_complete(self, success, message):
        """处理完成"""
        # 停止进度条
        self.progress.stop()

        # 启用按钮
        self.process_btn.config(state='normal', text="开始处理")

        if success:
            self.status_label.config(text="处理完成！", fg="green")
            if os.path.isfile(message):
                response = messagebox.askyesno("完成", "转换完成！是否打开文件？")
                if response:
                    os.startfile(message)
        else:
            self.status_label.config(text=f"处理失败: {message}", fg="red")
            messagebox.showerror("错误", message)


def main():
    root = tk.Tk()
    app = SimplePDFConverter(root)
    root.mainloop()


if __name__ == "__main__":
    main()